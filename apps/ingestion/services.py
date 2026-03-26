from __future__ import annotations

import hashlib
import json
import os
import re
import tempfile
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from django.conf import settings
from django.core.exceptions import SuspiciousFileOperation
from django.core.files.uploadedfile import UploadedFile
from django.utils import timezone
from docx import Document
from pypdf import PdfReader

from apps.publications.models import (
    Author,
    Keyword,
    Publication,
    PublicationChunk,
    PublicationLanguage,
    PublicationSubtype,
    PublicationType,
    ScientificSupervisor,
    TEXT_EXTRACTION_STATUS_CHOICES,
)
from apps.vector_store.services import VectorStoreService


@dataclass(slots=True)
class ExtractedSegment:
    text: str
    page_start: int | None = None
    page_end: int | None = None
    section_title: str = ""


@dataclass(slots=True)
class ChunkPayload:
    chunk_index: int
    text: str
    chunk_text: str
    source_kind: str = "fulltext"
    page_start: int | None = None
    page_end: int | None = None
    section_title: str = ""
    char_count: int = 0
    word_count: int = 0
    index_quality: float = 1.0


@dataclass(slots=True)
class FileExtractionAnalysis:
    file_extension: str = ""
    status: str = "pending"
    notes: str = ""
    supported_format: bool = False
    has_extractable_text: bool = False
    extracted_text: str = ""
    raw_text: str = ""
    segments: list[ExtractedSegment] = field(default_factory=list)
    page_count: int | None = None


_WORD_RE = re.compile(r"\S+")
_SENTENCE_SPLIT_RE = re.compile(r"(?<=[.!?])\s+|\n{2,}")
_WHITESPACE_RE = re.compile(r"\s+")
_YEAR_RE = re.compile(r"(?<!\d)(19\d{2}|20\d{2}|21\d{2})(?!\d)")
_TOKEN_RE = re.compile(r"[A-Za-zА-Яа-яЁё0-9-]{3,}")
_FILENAME_SEPARATORS_RE = re.compile(r"[_\-]+")

CYRILLIC_RE = re.compile(r"[А-Яа-яЁё]")
LATIN_RE = re.compile(r"[A-Za-z]")
_REFERENCE_HEADING_RE = re.compile(
    r"^(references|bibliography|список литературы|литература|bibliographic references)",
    re.IGNORECASE,
)
_REFERENCE_LINE_RE = re.compile(r"(?m)^\s*(?:\[?\d+[\].)]|\d+\.)\s+")
_REFERENCE_SIGNAL_RE = re.compile(
    r"(?:et\s+al\.?|doi\s*[:/]|vol\.?\s*\d+|pp?\.\s*\d+|j bone joint surg|clin orthop|bmj\b|cochrane\b)",
    re.IGNORECASE,
)
_DOT_LEADER_RE = re.compile(r"\.{5,}")
_TOC_HEADING_RE = re.compile(r"^(?:содержание|оглавление|table of contents)\b", re.IGNORECASE)
_TOC_LINE_RE = re.compile(r"(?m)^\s*[^\n]{3,180}?\.{5,}\s*\d{1,4}\s*$")
_PAGE_NUMBER_LINE_RE = re.compile(r"(?m)^\s*[\divxlcdmIVXLCDM\.\-–— ]{1,12}\s*$")
_SECTION_NUMBER_RE = re.compile(r"^(?:\d+(?:[\.\)]\d+)*[\.\)]?|[IVXLCDM]+[\.\)])\s+", re.IGNORECASE)
_HEADING_STYLE_RE = re.compile(r"heading|заголов", re.IGNORECASE)
_GENERIC_SECTION_TITLE_RE = re.compile(r"^(?:введение|заключение|выводы|содержание|оглавление|приложение|abstract|introduction|conclusion|references?)$", re.IGNORECASE)


def detect_script_kind(text: str) -> str:
    cyrillic = len(CYRILLIC_RE.findall(text or ""))
    latin = len(LATIN_RE.findall(text or ""))
    alpha = cyrillic + latin
    if alpha == 0:
        return "neutral"
    cyr_ratio = cyrillic / alpha
    lat_ratio = latin / alpha
    if cyr_ratio >= 0.7:
        return "cyrillic"
    if lat_ratio >= 0.7:
        return "latin"
    return "mixed"


def is_reference_heavy_text(text: str) -> bool:
    normalized = normalize_text(text)
    if not normalized:
        return False
    lowered = normalized.lower()
    if _REFERENCE_HEADING_RE.search(lowered):
        return True
    year_hits = len(_YEAR_RE.findall(normalized))
    numbered_refs = len(_REFERENCE_LINE_RE.findall(text or ""))
    signals = len(_REFERENCE_SIGNAL_RE.findall(lowered))
    punctuation_density = sum(1 for char in normalized if char in ';:[]()') / max(len(normalized), 1)
    return (
        numbered_refs >= 2
        or (year_hits >= 4 and signals >= 1)
        or (year_hits >= 6 and punctuation_density > 0.05)
    )


def is_table_of_contents_text(text: str) -> bool:
    normalized = normalize_text(text)
    if not normalized:
        return False

    lowered = normalized.lower()
    if _TOC_HEADING_RE.search(lowered):
        return True

    dot_leaders = len(_DOT_LEADER_RE.findall(text or ""))
    toc_lines = len(_TOC_LINE_RE.findall(text or ""))
    page_only_lines = len(_PAGE_NUMBER_LINE_RE.findall(text or ""))
    dot_ratio = sum(1 for char in (text or "") if char == ".") / max(len(text or ""), 1)

    return (
        toc_lines >= 2
        or (dot_leaders >= 3 and dot_ratio > 0.12)
        or (dot_leaders >= 2 and page_only_lines >= 1)
        or ("содержание" in lowered and dot_leaders >= 1)
    )


def chunk_index_quality(text: str) -> float:
    normalized = normalize_text(text)
    if not normalized:
        return 0.0
    alpha_count = sum(char.isalpha() for char in normalized)
    digit_count = sum(char.isdigit() for char in normalized)
    alpha_ratio = alpha_count / max(len(normalized), 1)
    digit_ratio = digit_count / max(len(normalized), 1)
    dot_ratio = sum(1 for char in (text or "") if char == ".") / max(len(text or ""), 1)
    info_density = chunk_information_density(normalized)
    quality = 1.0
    if alpha_ratio < 0.45:
        quality *= 0.65
    if digit_ratio > 0.18:
        quality *= 0.8
    if dot_ratio > 0.10:
        quality *= 0.45
    if info_density < float(getattr(settings, "VECTOR_CHUNK_MIN_INFORMATION_DENSITY", 0.24)):
        quality *= 0.62
    elif info_density > 0.58:
        quality *= 1.08
    if is_table_of_contents_text(text):
        quality *= 0.05
    if is_reference_heavy_text(text):
        quality *= 0.12
    return max(0.0, min(1.0, quality))


def _clip01(value: float) -> float:
    return max(0.0, min(1.0, float(value)))


def derive_publication_characteristics(
    publication: Publication,
    analysis: FileExtractionAnalysis,
    extracted_text: str = "",
) -> list[dict[str, Any]]:
    source_text = extracted_text or analysis.extracted_text or analysis.raw_text or publication.contents or ""
    normalized = normalize_text(source_text)
    lowered = normalized.lower()
    characteristics: list[dict[str, Any]] = []

    def add(slug: str, label: str, score: float, detail: str) -> None:
        clipped = _clip01(score)
        if clipped < 0.34:
            return
        characteristics.append({"slug": slug, "label": label, "score": round(clipped, 3), "detail": detail})

    token_count = max(len(_TOKEN_RE.findall(lowered)), 1)
    digit_ratio = sum(ch.isdigit() for ch in normalized) / max(len(normalized), 1)
    bullet_lines = len(BULLET_LINE_RE.findall(source_text or ""))
    table_lines = len(TABLE_LINE_RE.findall(source_text or ""))
    stats_hits = len(STATISTICAL_SIGNAL_RE.findall(lowered))
    instruction_hits = len(INSTRUCTION_SIGNAL_RE.findall(lowered))
    legal_hits = len(LEGAL_SIGNAL_RE.findall(lowered))
    educational_hits = len(EDUCATIONAL_SIGNAL_RE.findall(lowered))
    formula_hits = len(FORMULA_SIGNAL_RE.findall(source_text or ""))

    if analysis.status == "metadata_only_nontext":
        add(
            "nontext_structure",
            "Нетекстовая или сканированная структура",
            0.98,
            "В документе не обнаружен достаточный объём машиночитаемого основного текста.",
        )
    elif analysis.status == "metadata_only_unsupported":
        add(
            "metadata_only",
            "Поиск в основном по метаданным",
            0.92,
            "Формат файла не поддерживает полноценное извлечение текста, поэтому основой анализа служат метаданные.",
        )

    stats_score = max(digit_ratio * 5.0, min(stats_hits / 4.0, 1.0), min(table_lines / 3.0, 1.0))
    add("statistics", "Высокая доля статистических или табличных данных", stats_score, "В тексте много числовых значений, таблиц или статистических маркеров.")

    instruction_score = max(min(instruction_hits / 4.0, 1.0), min(bullet_lines / 4.0, 1.0))
    add("instructions", "Пошаговые инструкции или регламент", instruction_score, "В тексте много инструктивных формулировок, шагов, требований и маркированных пунктов.")

    legal_score = min(legal_hits / 4.0, 1.0)
    add("administrative", "Организационно-правовой или административный материал", legal_score, "Документ содержит лексику правил, сроков, выплат, заявок или договорных условий.")

    educational_score = min(educational_hits / 4.0, 1.0)
    add("educational", "Учебно-методический материал", educational_score, "В тексте присутствуют признаки методических рекомендаций, курса, заданий или учебной дисциплины.")

    formula_score = max(min(formula_hits / 4.0, 1.0), 0.0)
    add("formula_dense", "Формулы и специальные обозначения", formula_score, "В документе заметна доля формул, специальных символов или математических обозначений.")

    if is_reference_heavy_text(normalized):
        add("reference_apparatus", "Развитый научно-справочный аппарат", 0.78, "Документ содержит выраженный список литературы или плотные библиографические ссылки.")

    if analysis.page_count and analysis.page_count >= 80:
        add("long_read", "Крупный по объёму документ", min(analysis.page_count / 200.0, 1.0), f"Объём документа составляет около {analysis.page_count} страниц.")

    characteristics.sort(key=lambda item: (-item["score"], item["label"]))
    return characteristics[:5]


SUPPORTED_TEXT_EXTENSIONS = {".pdf", ".docx"}
STATUS_LABELS = dict(TEXT_EXTRACTION_STATUS_CHOICES)
RUSSIAN_STOPWORDS = {
    "это", "как", "или", "для", "при", "над", "под", "без", "что", "его", "её", "она", "они",
    "мы", "вы", "из", "по", "на", "до", "от", "за", "не", "но", "а", "и", "в", "во", "с",
    "со", "к", "ко", "у", "о", "об", "обо", "ли", "же", "бы", "быть", "так", "также", "если",
    "или", "для", "при", "между", "через", "после", "перед", "когда", "где", "который", "которая",
    "которые", "данный", "данная", "данные", "издание", "документ", "материал", "работа", "система",
}
ENGLISH_STOPWORDS = {
    "the", "and", "for", "with", "from", "that", "this", "into", "about", "document", "publication",
    "system", "work", "study", "article", "guide", "manual", "report", "using", "used", "are", "was",
}


STATISTICAL_SIGNAL_RE = re.compile(r"(?:табл\.?|table\b|рис\.?|figure\b|dataset|выборк|статистик|median|mean|sd\b|p-value|доверительн)", re.IGNORECASE)
INSTRUCTION_SIGNAL_RE = re.compile(r"(?:необходимо|следует|нужно|порядок|шаг(?:и)?|инструкц|получени[ея]|заявк|предостав(?:ить|ьте)|submit|must\b|should\b|instructions?)", re.IGNORECASE)
LEGAL_SIGNAL_RE = re.compile(r"(?:положение|регламент|приказ|договор|заявление|фонд|стипенд|выплат|реквизит|обязан|срок(?:и)?|конкурс)", re.IGNORECASE)
EDUCATIONAL_SIGNAL_RE = re.compile(r"(?:методическ|учебн|лабораторн|практическ|самостоятельн|курс|дисциплин|семинар|задани[ея]|рекомендац)", re.IGNORECASE)
FORMULA_SIGNAL_RE = re.compile(r"(?:формул|equation|lemma|theorem|proof|\b[a-zа-яё]\s*=\s*[^\s]+|[=±×÷∑∫λμσβγα])", re.IGNORECASE)
TABLE_LINE_RE = re.compile(r"(?m)^\s*[^\n]{0,120}\|[^\n]{0,120}$")
BULLET_LINE_RE = re.compile(r"(?m)^\s*(?:[-–—•●◦*]|\d+[\.)])\s+")


def normalize_text(value: str | None) -> str:
    return _WHITESPACE_RE.sub(" ", (value or "").strip())


def save_uploaded_file_to_temp(uploaded_file: UploadedFile) -> Path:
    temp_dir = Path(settings.MEDIA_ROOT) / "_tmp_upload_analysis"
    temp_dir.mkdir(parents=True, exist_ok=True)
    suffix = Path(uploaded_file.name or "").suffix
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix, dir=temp_dir) as temp_handle:
        for chunk in uploaded_file.chunks():
            temp_handle.write(chunk)
        return Path(temp_handle.name)


def _table_texts(document: Document) -> list[str]:
    texts: list[str] = []
    for table in document.tables:
        for row in table.rows:
            row_parts = [normalize_text(cell.text) for cell in row.cells if normalize_text(cell.text)]
            if row_parts:
                texts.append(" | ".join(row_parts))
    return texts


def chunk_information_density(text: str) -> float:
    normalized = normalize_text(text)
    if not normalized:
        return 0.0

    tokens = [token.lower() for token in _TOKEN_RE.findall(normalized)]
    if not tokens:
        return 0.0

    stopwords = RUSSIAN_STOPWORDS | ENGLISH_STOPWORDS
    informative = [token for token in tokens if token not in stopwords and not token.isdigit()]
    if not informative:
        return 0.0

    informative_ratio = len(informative) / len(tokens)
    unique_ratio = len(set(informative)) / len(informative)
    avg_token_len = sum(len(token) for token in informative) / len(informative)
    alpha_ratio = sum(char.isalpha() for char in normalized) / max(len(normalized), 1)
    density = (0.42 * informative_ratio) + (0.33 * unique_ratio) + (0.15 * min(avg_token_len / 8.0, 1.0)) + (0.10 * alpha_ratio)
    return max(0.0, min(1.0, density))


def is_heading_candidate(text: str, *, style_name: str = "") -> bool:
    normalized = normalize_text(text)
    if not normalized:
        return False
    if is_table_of_contents_text(normalized) or is_reference_heavy_text(normalized):
        return False

    style_name = normalize_text(style_name)
    if style_name and _HEADING_STYLE_RE.search(style_name):
        return True

    words = normalized.split()
    if len(words) == 0 or len(words) > int(getattr(settings, "VECTOR_HEADING_MAX_WORDS", 14)):
        return False

    if normalized.endswith((".", ";", ":")):
        return False

    digit_ratio = sum(char.isdigit() for char in normalized) / max(len(normalized), 1)
    punct_ratio = sum(1 for char in normalized if not char.isalnum() and not char.isspace()) / max(len(normalized), 1)
    upper_ratio = sum(1 for char in normalized if char.isalpha() and char.isupper()) / max(sum(1 for char in normalized if char.isalpha()), 1)

    if _SECTION_NUMBER_RE.match(normalized):
        return True
    if _GENERIC_SECTION_TITLE_RE.match(normalized):
        return True
    if upper_ratio > 0.72 and len(words) <= 10:
        return True
    if digit_ratio < 0.12 and punct_ratio < 0.10 and len(words) <= 8:
        titleish = sum(1 for word in words if word[:1].isupper()) >= max(1, len(words) - 1)
        if titleish:
            return True
    return False


def _lines_to_segments(lines: list[tuple[str, str]], *, page_number: int | None = None) -> list[ExtractedSegment]:
    segments: list[ExtractedSegment] = []
    current_heading = ""
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        paragraph_text = normalize_text(" ".join(paragraph_lines))
        if paragraph_text:
            segments.append(
                ExtractedSegment(
                    text=paragraph_text,
                    page_start=page_number,
                    page_end=page_number,
                    section_title=current_heading,
                )
            )
        paragraph_lines = []

    for raw_text, style_name in lines:
        normalized = normalize_text(raw_text)
        if not normalized:
            flush_paragraph()
            continue
        if is_heading_candidate(normalized, style_name=style_name):
            flush_paragraph()
            current_heading = normalized
            continue
        paragraph_lines.append(normalized)

    flush_paragraph()
    return segments


def _extract_pdf_segments(reader: PdfReader) -> tuple[list[ExtractedSegment], list[str]]:
    segments: list[ExtractedSegment] = []
    raw_pages: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        raw_page_text = (page.extract_text() or "").strip()
        if raw_page_text:
            raw_pages.append(raw_page_text)
        page_lines = [(line, "") for line in raw_page_text.splitlines()]
        page_segments = _lines_to_segments(page_lines, page_number=page_number)
        if page_segments:
            segments.extend(page_segments)
        else:
            normalized_page_text = normalize_text(raw_page_text)
            if normalized_page_text:
                segments.append(ExtractedSegment(text=normalized_page_text, page_start=page_number, page_end=page_number))
    return segments, raw_pages


def _extract_docx_segments(path: Path) -> tuple[list[ExtractedSegment], str]:
    document = Document(str(path))
    lines: list[tuple[str, str]] = []
    raw_lines: list[str] = []
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text or ""
        raw_lines.append(paragraph_text)
        lines.append((paragraph_text, getattr(getattr(paragraph, "style", None), "name", "")))
    for table_text in _table_texts(document):
        raw_lines.append(table_text)
        lines.append((table_text, ""))
    segments = _lines_to_segments(lines, page_number=None)
    raw_text = "\n".join(line for line in raw_lines if line)
    if not segments and raw_text.strip():
        normalized = normalize_text(raw_text)
        if normalized:
            segments = [ExtractedSegment(text=normalized)]
    return segments, raw_text


def _analyze_text_quality(text: str, *, page_count: int | None = None) -> tuple[bool, str]:
    normalized = normalize_text(text)
    char_threshold = int(getattr(settings, "INGESTION_MIN_TEXT_CHARS", 120))
    word_threshold = int(getattr(settings, "INGESTION_MIN_TEXT_WORDS", 25))
    char_count = len(normalized)
    word_count = len(_WORD_RE.findall(normalized))

    if char_count >= char_threshold and word_count >= word_threshold:
        return True, ""

    if page_count and page_count > 0:
        return (
            False,
            f"Извлечённого текста недостаточно для полноценной индексации ({word_count} слов на {page_count} стр.). "
            "Документ обработан в режиме metadata-only.",
        )
    return (
        False,
        f"Извлечённого текста недостаточно для полноценной индексации ({word_count} слов). "
        "Документ обработан в режиме metadata-only.",
    )


def analyze_publication_file(file_path: str | Path, *, original_name: str | None = None) -> FileExtractionAnalysis:
    path = Path(file_path)
    suffix = path.suffix.lower()
    analysis = FileExtractionAnalysis(file_extension=suffix)

    if not path.exists():
        analysis.status = "metadata_only_missing"
        analysis.notes = "Файл не найден в файловом хранилище. Для поиска будут использованы только метаданные."
        return analysis

    if suffix not in SUPPORTED_TEXT_EXTENSIONS:
        analysis.status = "metadata_only_unsupported"
        analysis.notes = (
            f"Формат {suffix or 'без расширения'} не поддерживается для автоматического извлечения текста. "
            "Для поиска будут использованы только введённые метаданные."
        )
        return analysis

    analysis.supported_format = True

    try:
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            analysis.page_count = len(reader.pages)
            segments, raw_pages = _extract_pdf_segments(reader)
            analysis.raw_text = "\n\n".join(raw_pages).strip()
            analysis.extracted_text = "\n".join(segment.text for segment in segments if segment.text).strip()
            analysis.segments = segments
        elif suffix == ".docx":
            segments, raw_text = _extract_docx_segments(path)
            analysis.raw_text = raw_text.strip()
            analysis.extracted_text = "\n\n".join(segment.text for segment in segments if segment.text).strip()
            analysis.segments = segments
    except Exception as exc:  # pragma: no cover - defensive fallback for malformed files
        analysis.status = "metadata_only_error"
        analysis.notes = (
            f"Не удалось автоматически извлечь текст ({exc.__class__.__name__}). "
            "Для поиска будут использованы только метаданные."
        )
        return analysis

    has_text, note = _analyze_text_quality(analysis.extracted_text, page_count=analysis.page_count)
    if has_text:
        analysis.status = "fulltext"
        analysis.has_extractable_text = True
        return analysis

    analysis.status = "metadata_only_nontext"
    analysis.notes = note or (
        "В файле не обнаружен машиночитаемый текст. Вероятно, документ имеет сканированную или нетекстовую структуру."
    )
    analysis.segments = []
    analysis.extracted_text = ""
    return analysis


def extract_text_from_file(file_path: str | Path) -> str:
    return analyze_publication_file(file_path).extracted_text


def extract_text_from_publication_file(publication: Publication) -> str:
    return analyze_publication_for_ingestion(publication).extracted_text


def extract_segments_from_file(file_path: str | Path) -> list[ExtractedSegment]:
    return analyze_publication_file(file_path).segments


def analyze_publication_for_ingestion(publication: Publication) -> FileExtractionAnalysis:
    if not publication.file:
        return FileExtractionAnalysis(
            file_extension="",
            status="metadata_only_missing",
            notes="Файл не приложен. Для поиска будут использованы только метаданные.",
        )
    try:
        file_name = publication.file.name
        if not file_name or not publication.file.storage.exists(file_name):
            return FileExtractionAnalysis(
                file_extension=Path(file_name or "").suffix.lower(),
                status="metadata_only_missing",
                notes="Файл отсутствует в файловом хранилище. Для поиска будут использованы только метаданные.",
            )
        return analyze_publication_file(publication.file.path, original_name=file_name)
    except (FileNotFoundError, OSError, SuspiciousFileOperation, ValueError):
        return FileExtractionAnalysis(
            file_extension=Path(getattr(publication.file, "name", "") or "").suffix.lower(),
            status="metadata_only_error",
            notes="Файл недоступен для чтения. Для поиска будут использованы только метаданные.",
        )


def _split_text_into_windows(text: str, max_words: int, overlap_words: int, max_chars: int) -> list[str]:
    normalized = normalize_text(text)
    if not normalized:
        return []

    units = [normalize_text(unit) for unit in _SENTENCE_SPLIT_RE.split(normalized) if normalize_text(unit)]
    if not units:
        units = [normalized]

    chunks: list[str] = []
    current_units: list[str] = []
    current_word_count = 0
    current_char_count = 0

    def flush() -> None:
        nonlocal current_units, current_word_count, current_char_count
        if not current_units:
            return
        chunk_text = normalize_text(" ".join(current_units))
        if chunk_text:
            chunks.append(chunk_text)
        overlap_seed: list[str] = []
        if overlap_words > 0 and current_units:
            words_remaining = overlap_words
            for unit in reversed(current_units):
                overlap_seed.insert(0, unit)
                words_remaining -= len(_WORD_RE.findall(unit))
                if words_remaining <= 0:
                    break
        current_units = overlap_seed
        current_word_count = sum(len(_WORD_RE.findall(unit)) for unit in current_units)
        current_char_count = sum(len(unit) for unit in current_units) + max(0, len(current_units) - 1)

    for unit in units:
        unit_words = len(_WORD_RE.findall(unit))
        unit_chars = len(unit)
        projected_words = current_word_count + unit_words
        projected_chars = current_char_count + unit_chars + (1 if current_units else 0)
        if current_units and (projected_words > max_words or projected_chars > max_chars):
            flush()
        if unit_words > max_words or unit_chars > max_chars:
            words = unit.split()
            step = max(1, max_words - overlap_words)
            for start in range(0, len(words), step):
                piece_words = words[start : start + max_words]
                piece_text = normalize_text(" ".join(piece_words))
                if piece_text:
                    chunks.append(piece_text[:max_chars].strip())
            current_units = []
            current_word_count = 0
            current_char_count = 0
            continue
        current_units.append(unit)
        current_word_count += unit_words
        current_char_count += unit_chars + (1 if len(current_units) > 1 else 0)

    if current_units:
        flush()

    deduped: list[str] = []
    previous = None
    for chunk in chunks:
        if chunk and chunk != previous:
            deduped.append(chunk)
            previous = chunk
    return deduped


def chunk_segments(segments: list[ExtractedSegment], max_words: int, overlap_words: int, max_chars: int) -> list[ChunkPayload]:
    chunk_payloads: list[ChunkPayload] = []
    chunk_index = 0

    for segment in segments:
        windows = _split_text_into_windows(segment.text, max_words=max_words, overlap_words=overlap_words, max_chars=max_chars)
        for window in windows:
            quality = chunk_index_quality(window)
            chunk_payloads.append(
                ChunkPayload(
                    chunk_index=chunk_index,
                    text=window,
                    chunk_text=window,
                    source_kind="fulltext",
                    page_start=segment.page_start,
                    page_end=segment.page_end,
                    section_title=segment.section_title,
                    char_count=len(window),
                    word_count=len(_WORD_RE.findall(window)),
                    index_quality=quality,
                )
            )
            chunk_index += 1
    return chunk_payloads


def _select_anchor_excerpt(text: str, *, max_chars: int) -> str:
    windows = _split_text_into_windows(text, max_words=80, overlap_words=0, max_chars=max_chars)
    if not windows:
        return ""
    ranked = sorted(
        windows,
        key=lambda item: (chunk_index_quality(item), chunk_information_density(item), -len(item)),
        reverse=True,
    )
    return ranked[0][:max_chars].strip()


def _build_metadata_anchor_chunk(publication: Publication, analysis: FileExtractionAnalysis, extracted_text: str) -> ChunkPayload | None:
    if not bool(getattr(settings, "VECTOR_INDEX_INCLUDE_METADATA_ANCHOR", True)):
        return None

    max_chars = int(getattr(settings, "VECTOR_ANCHOR_MAX_CHARS", 520))
    keyword_text = ", ".join(keyword.name for keyword in publication.keywords.all())
    authors = ", ".join(author.full_name for author in publication.authors.all())
    excerpt_source = publication.contents or extracted_text or analysis.raw_text or analysis.extracted_text
    excerpt = _select_anchor_excerpt(excerpt_source, max_chars=max_chars)
    parts = [publication.title]
    if publication.publication_subtype:
        parts.append(f"Подтип: {publication.publication_subtype.name}")
    if publication.language:
        parts.append(f"Язык: {publication.language.name}")
    if authors:
        parts.append(f"Авторы: {authors}")
    if keyword_text:
        parts.append(f"Ключевые слова: {keyword_text}")
    if excerpt:
        parts.append(f"Краткое содержание: {excerpt}")
    anchor_text = normalize_text("\n".join(part for part in parts if part))
    if not anchor_text:
        return None
    quality = max(0.72, chunk_index_quality(anchor_text))
    return ChunkPayload(
        chunk_index=-1,
        text=anchor_text[:max_chars],
        chunk_text=anchor_text[:max_chars],
        source_kind="metadata",
        section_title="Описание документа",
        char_count=len(anchor_text[:max_chars]),
        word_count=len(_WORD_RE.findall(anchor_text[:max_chars])),
        index_quality=quality,
    )


def build_search_document(publication: Publication, extracted_text: str = "") -> str:
    keyword_text = ", ".join(keyword.name for keyword in publication.keywords.all())
    authors = ", ".join(author.full_name for author in publication.authors.all())
    supervisors = ", ".join(supervisor.full_name for supervisor in publication.scientific_supervisors.all())
    publication_type = publication.publication_type.name if publication.publication_type else ""
    publication_subtype = publication.publication_subtype.name if publication.publication_subtype else ""
    publishers = ", ".join(publisher.name for publisher in publication.publishers.all())
    places = ", ".join(place.name for place in publication.publication_places.all())
    parts = [
        publication.title,
        authors,
        supervisors,
        publication_type,
        publication_subtype,
        publication.contents,
        publication.grif_text,
        publication.grant_text,
        keyword_text,
        publishers,
        places,
        extracted_text,
    ]
    return "\n\n".join(part for part in parts if part).strip()


def build_chunk_vector_document(
    publication: Publication,
    chunk_text: str,
    source_kind: str = "fulltext",
    section_title: str = "",
) -> str:
    header = _build_chunk_context_header(publication)
    prefix = "Фрагмент документа" if source_kind == "fulltext" else "Метаданные документа"
    body_parts: list[str] = []
    if section_title:
        body_parts.append(f"Раздел: {section_title}")
    body_parts.append(f"{prefix}: {chunk_text}")
    body = "\n".join(part for part in body_parts if part)
    return normalize_text(f"{header}\n\n{body}" if header else body)

def compute_publication_index_signature(publication: Publication) -> str:
    file_info: dict[str, str | int | None] = {
        "name": getattr(publication.file, "name", "") or "",
        "exists": False,
        "size": None,
        "mtime_ns": None,
    }
    try:
        file_name = getattr(publication.file, "name", "") or ""
        storage = getattr(publication.file, "storage", None)
        if file_name and storage and storage.exists(file_name):
            file_info["exists"] = True
            try:
                file_info["size"] = int(storage.size(file_name))
            except Exception:
                file_info["size"] = None
            try:
                file_path = publication.file.path
                stat = Path(file_path).stat()
                file_info["mtime_ns"] = int(stat.st_mtime_ns)
            except Exception:
                file_info["mtime_ns"] = None
    except Exception:
        pass

    payload = {
        "schema_version": getattr(settings, "VECTOR_INDEX_SCHEMA_VERSION", "v1"),
        "model": getattr(settings, "MILVUS_BGE_M3_MODEL", "BAAI/bge-m3"),
        "chunk_words": int(getattr(settings, "VECTOR_CHUNK_MAX_WORDS", 320)),
        "chunk_overlap": int(getattr(settings, "VECTOR_CHUNK_OVERLAP_WORDS", 40)),
        "chunk_chars": int(getattr(settings, "VECTOR_CHUNK_MAX_CHARS", 2200)),
        "min_text_chars": int(getattr(settings, "INGESTION_MIN_TEXT_CHARS", 120)),
        "min_text_words": int(getattr(settings, "INGESTION_MIN_TEXT_WORDS", 25)),
        "min_chunk_quality": float(getattr(settings, "VECTOR_CHUNK_MIN_INDEX_QUALITY", 0.28)),
        "min_information_density": float(getattr(settings, "VECTOR_CHUNK_MIN_INFORMATION_DENSITY", 0.24)),
        "include_metadata_anchor": bool(getattr(settings, "VECTOR_INDEX_INCLUDE_METADATA_ANCHOR", True)),
        "anchor_max_chars": int(getattr(settings, "VECTOR_ANCHOR_MAX_CHARS", 520)),
        "heading_max_words": int(getattr(settings, "VECTOR_HEADING_MAX_WORDS", 14)),
        "title": publication.title,
        "contents": publication.contents,
        "grif_text": publication.grif_text,
        "grant_text": publication.grant_text,
        "publication_year": publication.publication_year,
        "language": publication.language.name if publication.language else "",
        "publication_type": publication.publication_type.name if publication.publication_type else "",
        "publication_subtype": publication.publication_subtype.name if publication.publication_subtype else "",
        "authors": [author.full_name for author in publication.authors.all()],
        "keywords": [keyword.name for keyword in publication.keywords.all()],
        "publishers": [publisher.name for publisher in publication.publishers.all()],
        "places": [place.name for place in publication.publication_places.all()],
        "supervisors": [supervisor.full_name for supervisor in publication.scientific_supervisors.all()],
        "file": file_info,
    }
    serialized = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(serialized.encode("utf-8")).hexdigest()


def _build_chunk_context_header(publication: Publication) -> str:
    metadata_parts = [publication.title]
    authors = ", ".join(author.full_name for author in publication.authors.all())
    if authors:
        metadata_parts.append(f"Авторы: {authors}")
    publication_type = publication.publication_type.name if publication.publication_type else ""
    if publication_type:
        metadata_parts.append(f"Тип: {publication_type}")
    if publication.publication_subtype:
        metadata_parts.append(f"Подтип: {publication.publication_subtype.name}")
    if publication.language:
        metadata_parts.append(f"Язык: {publication.language.name}")
    keywords = ", ".join(keyword.name for keyword in publication.keywords.all())
    if keywords:
        metadata_parts.append(f"Ключевые слова: {keywords}")
    if publication.publication_year:
        metadata_parts.append(f"Год: {publication.publication_year}")
    return "\n".join(part for part in metadata_parts if part)


def build_publication_chunks(
    publication: Publication,
    analysis: FileExtractionAnalysis | None = None,
    extracted_text: str = "",
) -> list[ChunkPayload]:
    max_words = int(getattr(settings, "VECTOR_CHUNK_MAX_WORDS", 320))
    overlap_words = int(getattr(settings, "VECTOR_CHUNK_OVERLAP_WORDS", 40))
    max_chars = int(getattr(settings, "VECTOR_CHUNK_MAX_CHARS", 2200))

    analysis = analysis or analyze_publication_for_ingestion(publication)
    segments = analysis.segments if analysis and analysis.has_extractable_text else []
    fulltext = extracted_text or analysis.extracted_text

    raw_chunks = chunk_segments(segments, max_words=max_words, overlap_words=overlap_words, max_chars=max_chars)
    min_quality = float(getattr(settings, "VECTOR_CHUNK_MIN_INDEX_QUALITY", 0.28))
    filtered_chunks = [chunk for chunk in raw_chunks if chunk.index_quality >= min_quality]
    raw_chunks = filtered_chunks or raw_chunks

    anchor_chunk = _build_metadata_anchor_chunk(publication, analysis, fulltext)
    if anchor_chunk is not None:
        raw_chunks = [anchor_chunk, *raw_chunks]

    if not raw_chunks:
        combined = build_search_document(publication, extracted_text=fulltext)
        combined_chunks = _split_text_into_windows(combined, max_words=max_words, overlap_words=overlap_words, max_chars=max_chars) if combined else []
        if not combined_chunks and combined:
            combined_chunks = [combined[:max_chars]]
        raw_chunks = [
            ChunkPayload(
                chunk_index=index,
                text=chunk_text,
                chunk_text=chunk_text,
                source_kind="metadata",
                section_title="Описание документа",
                char_count=len(chunk_text),
                word_count=len(_WORD_RE.findall(chunk_text)),
                index_quality=chunk_index_quality(chunk_text),
            )
            for index, chunk_text in enumerate(combined_chunks)
            if chunk_text
        ]
    payloads: list[ChunkPayload] = []
    for raw_chunk in raw_chunks:
        contextual_text = build_chunk_vector_document(
            publication,
            raw_chunk.text,
            raw_chunk.source_kind,
            section_title=raw_chunk.section_title,
        )
        payloads.append(
            ChunkPayload(
                chunk_index=raw_chunk.chunk_index,
                text=raw_chunk.text,
                chunk_text=contextual_text,
                source_kind=raw_chunk.source_kind,
                page_start=raw_chunk.page_start,
                page_end=raw_chunk.page_end,
                section_title=raw_chunk.section_title,
                char_count=raw_chunk.char_count,
                word_count=raw_chunk.word_count,
                index_quality=raw_chunk.index_quality,
            )
        )
    return payloads


def sync_publication_chunks(publication: Publication, chunk_payloads: list[ChunkPayload]) -> list[PublicationChunk]:
    publication.chunks.all().delete()
    normalized_payloads: list[ChunkPayload] = []
    for normalized_index, payload in enumerate(chunk_payloads):
        if not payload.text:
            continue
        payload.chunk_index = normalized_index
        normalized_payloads.append(payload)
    chunk_objects = [
        PublicationChunk(
            publication=publication,
            chunk_index=payload.chunk_index,
            text=payload.text,
            source_kind=payload.source_kind,
            page_start=payload.page_start,
            page_end=payload.page_end,
            section_title=payload.section_title,
            char_count=payload.char_count,
            word_count=payload.word_count,
            index_quality=payload.index_quality,
        )
        for payload in normalized_payloads
    ]
    if not chunk_objects:
        return []
    created_objects = PublicationChunk.objects.bulk_create(chunk_objects)
    return list(created_objects)


def get_existing_publication_chunks(publication: Publication) -> list[PublicationChunk]:
    chunks = list(publication.chunks.all().order_by("chunk_index"))
    for chunk in chunks:
        chunk.publication = publication
    return chunks


def rebuild_publication_chunks(publication: Publication) -> tuple[FileExtractionAnalysis, list[PublicationChunk]]:
    analysis = analyze_publication_for_ingestion(publication)
    update_fields = ["file_extension", "text_extraction_status", "text_extraction_notes", "has_extracted_text", "derived_characteristics"]
    publication.file_extension = analysis.file_extension
    publication.text_extraction_status = analysis.status
    publication.text_extraction_notes = analysis.notes
    publication.has_extracted_text = analysis.has_extractable_text
    publication.derived_characteristics = derive_publication_characteristics(publication, analysis, analysis.extracted_text)
    publication.save(update_fields=update_fields)

    chunk_payloads = build_publication_chunks(publication, analysis=analysis, extracted_text=analysis.extracted_text)
    chunk_objects = sync_publication_chunks(publication, chunk_payloads)
    for chunk in chunk_objects:
        chunk.publication = publication
    return analysis, chunk_objects


def ingest_publication(publication: Publication, index_in_vector_store: bool = True) -> Publication:
    _, chunk_objects = rebuild_publication_chunks(publication)
    signature = compute_publication_index_signature(publication)

    if index_in_vector_store and not publication.is_draft:
        service = VectorStoreService()
        service.replace_publication_chunks(publication=publication, chunks=chunk_objects)
        publication.vector_indexed_at = timezone.now()
        publication.vector_index_signature = signature
        publication.save(update_fields=["vector_indexed_at", "vector_index_signature"])
    return publication


# --- Metadata suggestion helpers -------------------------------------------------

def _normalize_for_match(text: str) -> str:
    return normalize_text(text).lower().replace("ё", "е")


def _candidate_title_lines(raw_text: str) -> list[str]:
    return [line.strip() for line in raw_text.splitlines() if line and line.strip()]


def _score_title_candidate(candidate: str) -> float:
    normalized = normalize_text(candidate)
    if not normalized:
        return -1.0
    score = 0.0
    words = normalized.split()
    word_count = len(words)
    if 2 <= word_count <= 18:
        score += 2.0
    if 12 <= len(normalized) <= 180:
        score += 2.0
    if sum(ch.isalpha() for ch in normalized) / max(len(normalized), 1) > 0.6:
        score += 1.5
    if not normalized.endswith(('.', ':', ';')):
        score += 0.5
    if normalized.isupper():
        score -= 0.3
    if normalized.lower().startswith(("министерство", "федеральное", "университет", "кафедра")):
        score -= 0.8
    if _YEAR_RE.search(normalized):
        score -= 0.5
    return score


def _humanize_filename(filename: str) -> str:
    stem = Path(filename).stem
    stem = _FILENAME_SEPARATORS_RE.sub(" ", stem)
    stem = re.sub(r"\s+", " ", stem).strip()
    return stem


def suggest_title_from_text(raw_text: str, filename: str = "") -> str:
    candidates = _candidate_title_lines(raw_text)[:25]
    if candidates:
        best = max(candidates, key=_score_title_candidate)
        if _score_title_candidate(best) >= 2.5:
            return normalize_text(best)
    humanized = _humanize_filename(filename)
    return normalize_text(humanized)


def suggest_year_from_text(raw_text: str, filename: str = "") -> int | None:
    matches = [int(value) for value in _YEAR_RE.findall(f"{filename}\n{raw_text}")]
    if not matches:
        return None
    current_year = 2100
    plausible = [value for value in matches if 1900 <= value <= current_year]
    if not plausible:
        return None
    counts = Counter(plausible)
    return counts.most_common(1)[0][0]


def suggest_language_from_text(text: str) -> PublicationLanguage | None:
    normalized = text or ""
    cyrillic_count = sum(1 for char in normalized if "а" <= char.lower() <= "я" or char.lower() == "ё")
    latin_count = sum(1 for char in normalized if "a" <= char.lower() <= "z")
    if cyrillic_count == 0 and latin_count == 0:
        return None
    target = "рус" if cyrillic_count >= latin_count else "англ"
    return PublicationLanguage.objects.filter(name__icontains=target).order_by("name").first()


def _find_existing_people_matches(text: str, model, field_name: str, limit: int = 5) -> list[int]:
    normalized_text = _normalize_for_match(text)
    matched_ids: list[int] = []
    for obj in model.objects.all().only("id", field_name):
        candidate = _normalize_for_match(getattr(obj, field_name, ""))
        if candidate and candidate in normalized_text:
            matched_ids.append(obj.pk)
        if len(matched_ids) >= limit:
            break
    return matched_ids


def _find_keyword_matches(text: str, limit: int = 8) -> list[int]:
    normalized_text = _normalize_for_match(text)
    matched_ids: list[int] = []
    for keyword in Keyword.objects.all().only("id", "name"):
        candidate = _normalize_for_match(keyword.name)
        if candidate and candidate in normalized_text:
            matched_ids.append(keyword.pk)
        if len(matched_ids) >= limit:
            break
    return matched_ids


def _pick_best_subtype(text: str) -> PublicationSubtype | None:
    normalized_text = _normalize_for_match(text)
    best_subtype = None
    best_score = 0
    for subtype in PublicationSubtype.objects.select_related("publication_type").all():
        name = _normalize_for_match(subtype.name)
        if not name:
            continue
        tokens = [token for token in re.split(r"\W+", name) if len(token) >= 4]
        if not tokens:
            tokens = [name]
        score = sum(1 for token in tokens if token in normalized_text)
        if name in normalized_text:
            score += 2
        if score > best_score:
            best_score = score
            best_subtype = subtype
    return best_subtype if best_score > 0 else None


def _pick_best_type(text: str) -> PublicationType | None:
    normalized_text = _normalize_for_match(text)
    best_type = None
    best_score = 0
    for publication_type in PublicationType.objects.all():
        name = _normalize_for_match(publication_type.name)
        if not name:
            continue
        tokens = [token for token in re.split(r"\W+", name) if len(token) >= 4]
        if not tokens:
            tokens = [name]
        score = sum(1 for token in tokens if token in normalized_text)
        if name in normalized_text:
            score += 1
        if score > best_score:
            best_score = score
            best_type = publication_type
    return best_type if best_score > 0 else None


def _suggest_contents_preview(text: str) -> str:
    normalized = normalize_text(text)
    if not normalized:
        return ""
    max_chars = int(getattr(settings, "UPLOAD_PREFILL_CONTENTS_MAX_CHARS", 700))
    if len(normalized) <= max_chars:
        return normalized
    return normalized[: max_chars - 1].rstrip() + "…"


def _suggest_keywords_from_frequency(text: str, limit: int = 8) -> list[str]:
    normalized = _normalize_for_match(text)
    stopwords = RUSSIAN_STOPWORDS | ENGLISH_STOPWORDS
    counts = Counter(
        token
        for token in _TOKEN_RE.findall(normalized)
        if token not in stopwords and not token.isdigit() and len(token) >= 4
    )
    return [token for token, _ in counts.most_common(limit)]


def generate_metadata_prefill(path: str | Path, filename: str = "") -> dict[str, Any]:
    analysis = analyze_publication_file(path, original_name=filename)
    search_text = analysis.raw_text or analysis.extracted_text or _humanize_filename(filename)
    title = suggest_title_from_text(analysis.raw_text or analysis.extracted_text, filename=filename)
    year = suggest_year_from_text(search_text, filename=filename)
    language = suggest_language_from_text(search_text)
    subtype = _pick_best_subtype(f"{title}\n{search_text}")
    publication_type = subtype.publication_type if subtype else _pick_best_type(f"{title}\n{search_text}")

    author_ids = _find_existing_people_matches(search_text, Author, "full_name")
    supervisor_ids = _find_existing_people_matches(
        search_text,
        ScientificSupervisor,
        "full_name",
    )
    keyword_ids = _find_keyword_matches(f"{title}\n{search_text}")
    keyword_names = [
        keyword.name
        for keyword in Keyword.objects.filter(pk__in=keyword_ids)
    ]
    if not keyword_names:
        keyword_names = _suggest_keywords_from_frequency(f"{title}\n{search_text}")

    selected_field_names = [
        name
        for name, value in {
            "title": title,
            "publication_year": year,
            "language": getattr(language, "pk", None),
            "publication_subtype": getattr(subtype, "pk", None),
            "authors": author_ids,
            "scientific_supervisors": supervisor_ids,
            "keywords": keyword_ids,
            "contents": _suggest_contents_preview(search_text),
        }.items()
        if value not in (None, "", [])
    ]

    characteristics = derive_publication_characteristics(
        publication=Publication(title=title or _humanize_filename(filename), contents=_suggest_contents_preview(search_text)),
        analysis=analysis,
        extracted_text=analysis.extracted_text,
    )

    return {
        "status": analysis.status,
        "status_label": STATUS_LABELS.get(analysis.status, analysis.status),
        "notes": analysis.notes,
        "supported_format": analysis.supported_format,
        "has_extractable_text": analysis.has_extractable_text,
        "file_extension": analysis.file_extension,
        "page_count": analysis.page_count,
        "title_candidate": title,
        "keyword_name_suggestions": keyword_names,
        "characteristics": characteristics,
        "selected_field_names": selected_field_names,
        "field_values": {
            "title": title,
            "publication_year": year,
            "language": getattr(language, "pk", None),
            "publication_subtype": getattr(subtype, "pk", None),
            "publication_type_label": getattr(publication_type, "name", ""),
            "authors": author_ids,
            "scientific_supervisors": supervisor_ids,
            "keywords": keyword_ids,
            "contents": _suggest_contents_preview(search_text),
        },
    }


def generate_metadata_prefill_from_upload(uploaded_file: UploadedFile) -> dict[str, Any]:
    temp_path = save_uploaded_file_to_temp(uploaded_file)
    try:
        return generate_metadata_prefill(temp_path, filename=uploaded_file.name or "")
    finally:
        try:
            os.unlink(temp_path)
        except OSError:
            pass
